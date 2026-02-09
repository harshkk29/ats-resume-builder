"""
Resume Generator Module
Creates professional PDF and DOCX resumes
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
import os

class ResumeGenerator:
    """Generate professional resumes in PDF and DOCX formats"""
    
    @staticmethod
    def generate_pdf(resume_data: Dict, output_path: str) -> bool:
        """
        Generate PDF resume
        
        Args:
            resume_data: Dictionary containing resume information
            output_path: Path to save PDF file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch,
                                   leftMargin=0.75*inch, rightMargin=0.75*inch)
            
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=6,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold',
                borderWidth=1,
                borderColor=colors.HexColor('#3498db'),
                borderPadding=5,
                backColor=colors.HexColor('#ecf0f1')
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=6,
                fontName='Helvetica'
            )
            
            bullet_style = ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                leftIndent=20,
                spaceAfter=4,
                fontName='Helvetica'
            )
            
            # Name
            name = resume_data.get('name', 'N/A')
            story.append(Paragraph(name, title_style))
            
            # Contact Info
            contact_parts = []
            if resume_data.get('email'):
                contact_parts.append(resume_data['email'])
            if resume_data.get('phone'):
                contact_parts.append(resume_data['phone'])
            
            contact_text = ' | '.join(contact_parts)
            contact_style = ParagraphStyle('Contact', parent=normal_style, alignment=TA_CENTER, fontSize=10)
            story.append(Paragraph(contact_text, contact_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
                story.append(Paragraph(resume_data['professional_summary'], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Skills
            if resume_data.get('skills'):
                story.append(Paragraph('SKILLS', heading_style))
                skills = resume_data['skills']
                
                if isinstance(skills, dict):
                    if skills.get('technical'):
                        tech_skills = ', '.join(skills['technical'])
                        story.append(Paragraph(f'<b>Technical:</b> {tech_skills}', normal_style))
                    if skills.get('soft'):
                        soft_skills = ', '.join(skills['soft'])
                        story.append(Paragraph(f'<b>Soft Skills:</b> {soft_skills}', normal_style))
                elif isinstance(skills, list):
                    skills_text = ', '.join(skills)
                    story.append(Paragraph(skills_text, normal_style))
                
                story.append(Spacer(1, 0.1*inch))
            
            # Experience
            if resume_data.get('experience'):
                story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
                for exp in resume_data['experience']:
                    title = exp.get('title', 'N/A')
                    company = exp.get('company', 'N/A')
                    duration = exp.get('duration', 'N/A')
                    
                    exp_header = f'<b>{title}</b> | {company} | {duration}'
                    story.append(Paragraph(exp_header, normal_style))
                    
                    responsibilities = exp.get('responsibilities', exp.get('achievements', []))
                    for resp in responsibilities:
                        story.append(Paragraph(f'• {resp}', bullet_style))
                    
                    story.append(Spacer(1, 0.1*inch))
            
            # Projects
            if resume_data.get('projects'):
                story.append(Paragraph('PROJECTS', heading_style))
                for proj in resume_data['projects']:
                    name = proj.get('name', 'N/A')
                    description = proj.get('description', '')
                    technologies = proj.get('technologies', [])
                    
                    proj_header = f'<b>{name}</b>'
                    if technologies:
                        tech_str = ', '.join(technologies)
                        proj_header += f' | <i>{tech_str}</i>'
                    
                    story.append(Paragraph(proj_header, normal_style))
                    if description:
                        story.append(Paragraph(description, bullet_style))
                    
                    achievements = proj.get('achievements', [])
                    for ach in achievements:
                        story.append(Paragraph(f'• {ach}', bullet_style))
                    
                    story.append(Spacer(1, 0.1*inch))
            
            # Education
            if resume_data.get('education'):
                story.append(Paragraph('EDUCATION', heading_style))
                for edu in resume_data['education']:
                    degree = edu.get('degree', 'N/A')
                    institution = edu.get('institution', 'N/A')
                    year = edu.get('year', 'N/A')
                    gpa = edu.get('gpa', '')
                    
                    edu_text = f'<b>{degree}</b> | {institution} | {year}'
                    if gpa:
                        edu_text += f' | GPA: {gpa}'
                    
                    story.append(Paragraph(edu_text, normal_style))
                    
                    coursework = edu.get('relevant_coursework', [])
                    if coursework:
                        coursework_text = 'Relevant Coursework: ' + ', '.join(coursework)
                        story.append(Paragraph(coursework_text, bullet_style))
                    
                    story.append(Spacer(1, 0.05*inch))
            
            # Certifications
            if resume_data.get('certifications'):
                story.append(Paragraph('CERTIFICATIONS', heading_style))
                certs = resume_data['certifications']
                for cert in certs:
                    story.append(Paragraph(f'• {cert}', bullet_style))
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return False
    
    @staticmethod
    def generate_docx(resume_data: Dict, output_path: str) -> bool:
        """
        Generate DOCX resume
        
        Args:
            resume_data: Dictionary containing resume information
            output_path: Path to save DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Name
            name = doc.add_paragraph()
            name_run = name.add_run(resume_data.get('name', 'N/A'))
            name_run.font.size = Pt(20)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(26, 26, 26)
            name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact Info
            contact = doc.add_paragraph()
            contact_parts = []
            if resume_data.get('email'):
                contact_parts.append(resume_data['email'])
            if resume_data.get('phone'):
                contact_parts.append(resume_data['phone'])
            
            contact_text = ' | '.join(contact_parts)
            contact_run = contact.add_run(contact_text)
            contact_run.font.size = Pt(11)
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacer
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                doc.add_paragraph(resume_data['professional_summary'])
            
            # Skills
            if resume_data.get('skills'):
                heading = doc.add_heading('SKILLS', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                
                skills = resume_data['skills']
                if isinstance(skills, dict):
                    if skills.get('technical'):
                        p = doc.add_paragraph()
                        p.add_run('Technical: ').bold = True
                        p.add_run(', '.join(skills['technical']))
                    if skills.get('soft'):
                        p = doc.add_paragraph()
                        p.add_run('Soft Skills: ').bold = True
                        p.add_run(', '.join(skills['soft']))
                elif isinstance(skills, list):
                    doc.add_paragraph(', '.join(skills))
            
            # Experience
            if resume_data.get('experience'):
                heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                
                for exp in resume_data['experience']:
                    p = doc.add_paragraph()
                    p.add_run(exp.get('title', 'N/A')).bold = True
                    p.add_run(f" | {exp.get('company', 'N/A')} | {exp.get('duration', 'N/A')}")
                    
                    responsibilities = exp.get('responsibilities', exp.get('achievements', []))
                    for resp in responsibilities:
                        doc.add_paragraph(resp, style='List Bullet')
            
            # Projects
            if resume_data.get('projects'):
                heading = doc.add_heading('PROJECTS', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                
                for proj in resume_data['projects']:
                    p = doc.add_paragraph()
                    p.add_run(proj.get('name', 'N/A')).bold = True
                    
                    technologies = proj.get('technologies', [])
                    if technologies:
                        tech_run = p.add_run(f" | {', '.join(technologies)}")
                        tech_run.italic = True
                    
                    if proj.get('description'):
                        doc.add_paragraph(proj['description'])
                    
                    achievements = proj.get('achievements', [])
                    for ach in achievements:
                        doc.add_paragraph(ach, style='List Bullet')
            
            # Education
            if resume_data.get('education'):
                heading = doc.add_heading('EDUCATION', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                
                for edu in resume_data['education']:
                    p = doc.add_paragraph()
                    p.add_run(edu.get('degree', 'N/A')).bold = True
                    p.add_run(f" | {edu.get('institution', 'N/A')} | {edu.get('year', 'N/A')}")
                    
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")
                    
                    coursework = edu.get('relevant_coursework', [])
                    if coursework:
                        cw_p = doc.add_paragraph()
                        cw_p.add_run('Relevant Coursework: ').italic = True
                        cw_p.add_run(', '.join(coursework))
            
            # Certifications
            if resume_data.get('certifications'):
                heading = doc.add_heading('CERTIFICATIONS', level=2)
                heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                
                for cert in resume_data['certifications']:
                    doc.add_paragraph(cert, style='List Bullet')
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error generating DOCX: {e}")
            return False
